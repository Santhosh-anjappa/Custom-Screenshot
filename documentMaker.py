from docx import Document
from docx.shared import Inches


class Doc():
    def __init__(self):
        self.myDoc = Document()

        # Set Left and Right Margin
        self.section = self.myDoc.sections[0]
        self.section.left_margin = Inches(0.5)
        self.section.right_margin = Inches(0.5)

        # Add a Table
        self.table = self.myDoc.add_table(rows = 0, cols = 0)
        self.table.style = self.myDoc.styles['TableGrid']
        self.table.autofit = False

        # Set Columns
        self.addColumns(3,[0.3,2.0,5.5])
        

    def addColumns(self, n=1, widthList=[]):
        if( n == len(widthList)):
            for i in range(n):
                self.table.add_column(Inches(widthList[i]))

    
    def addRows(self,n=1):
        # for i in range(n):
        #     self.table.add_row()
        return self.table.add_row().cells

    def addImageToCell(self, pic, row_cells):
        paragraph = row_cells[2].paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(pic, width = Inches(5.0))
        row_cells[0].add_paragraph(pic[-5]+'.')

    def addComment(self, comment,rowCells):
        rowCells[1].add_paragraph(comment)

    def save(self, Path):
        self.myDoc.save(Path)


if __name__ == "__main__":
    doc = Doc()
    doc.addRows(2)
    doc.save('C:\\Users\\736131\\Documents\\test2.docx')